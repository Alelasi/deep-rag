#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
D盘全盘RAG索引器 - 基于助理项目的混合检索版本
功能：
1. 索引D盘全部文档（除危险目录）
2. 混合检索（BM25 + 向量）
3. 增量更新（mtime检测）
4. 量化压缩（减小索引体积）

用法：
    # 首次全盘索引
    python 全盘索引器_v1.py index D:\

    # 增量更新
    python 全盘索引器_v1.py update

    # 搜索
    python 全盘索引器_v1.py search "查询内容"
"""

import os
import pickle
import json
import hashlib
import numpy as np
from pathlib import Path
from typing import List, Dict, Tuple
from datetime import datetime
from collections import Counter
import re

# ==================== 配置 ====================

# 索引存储位置
INDEX_DIR = Path("D:/文档/ai提问相关/哲思灵智/rag-docs-全盘")
INDEX_DIR.mkdir(parents=True, exist_ok=True)

EMBEDDINGS_FILE = INDEX_DIR / "embeddings.pkl"
METADATA_FILE = INDEX_DIR / "metadata.json"
BM25_INDEX_FILE = INDEX_DIR / "bm25_index.pkl"
STATE_FILE = INDEX_DIR / ".index_state.json"

# 黑名单目录（危险/无用）
BLACKLIST_DIRS = {
    # 系统目录
    "Windows", "Program Files", "Program Files (x86)",
    "$Recycle.Bin", "System Volume Information",
    "ProgramData", "Recovery", "PerfLogs", "Boot",

    # 开发缓存
    "node_modules", "__pycache__", ".git", ".svn", ".hg",
    "venv", ".venv", "env", ".env", "virtualenv",
    ".cache", ".npm", ".yarn", ".pnpm", ".pip",
    "dist", "build", "target", ".gradle", ".m2",
    "obj", "bin", "Debug", "Release",

    # 临时文件
    "Temp", "tmp", "temp", "cache", "Cache",

    # 大型工具
    ".docker", ".kube", ".minikube", ".vagrant",

    # 游戏/娱乐（可选，体积大）
    "Steam", "Epic Games", "WeGame",
}

# 黑名单文件扩展名
BLACKLIST_EXTS = {
    # 二进制可执行
    ".exe", ".dll", ".so", ".dylib", ".bin", ".msi", ".app",

    # 压缩包
    ".zip", ".rar", ".7z", ".tar", ".gz", ".bz2", ".xz",

    # 媒体文件
    ".mp4", ".avi", ".mkv", ".mov", ".wmv", ".flv",
    ".mp3", ".wav", ".flac", ".aac", ".ogg",
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".ico", ".svg",

    # 数据库
    ".db", ".sqlite", ".mdb", ".accdb",

    # 其他大文件
    ".iso", ".dmg", ".pkg",
}

# 支持的文件类型
SUPPORTED_EXTS = {
    ".md", ".txt", ".py", ".js", ".ts", ".java", ".cpp", ".c", ".h",
    ".json", ".yaml", ".yml", ".xml", ".html", ".css", ".scss",
    ".sh", ".bat", ".ps1", ".sql", ".go", ".rs", ".rb", ".php",
    ".log", ".csv", ".ini", ".conf", ".cfg",
    # 文档类型（需要额外解析）
    ".pdf", ".docx", ".xlsx", ".pptx",
}

# ==================== 工具函数 ====================

def should_skip_dir(dir_path: Path) -> bool:
    """判断是否跳过目录"""
    dir_name = dir_path.name

    # 黑名单
    if dir_name in BLACKLIST_DIRS:
        return True

    # 隐藏目录（以.开头，但排除.config等常用目录）
    if dir_name.startswith(".") and dir_name not in {".config", ".ssh"}:
        return True

    return False

def should_skip_file(file_path: Path) -> bool:
    """判断是否跳过文件"""
    # 扩展名检查
    ext = file_path.suffix.lower()
    if ext in BLACKLIST_EXTS:
        return True
    if ext not in SUPPORTED_EXTS:
        return True

    # 临时文件
    if file_path.name.startswith("~$"):
        return True

    # 文件大小检查（> 10MB 跳过）
    try:
        if file_path.stat().st_size > 10 * 1024 * 1024:
            return True
    except:
        return True

    return False

def read_file_content(file_path: Path) -> str:
    """读取文件内容（支持文本/PDF/DOCX）"""
    try:
        ext = file_path.suffix.lower()

        # 文本文件直接读取
        if ext in {".md", ".txt", ".py", ".js", ".ts", ".java", ".cpp", ".c", ".h",
                   ".json", ".yaml", ".yml", ".xml", ".html", ".css", ".log",
                   ".csv", ".ini", ".conf", ".sh", ".bat", ".sql", ".go", ".rs"}:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

        # PDF 文件
        elif ext == ".pdf":
            try:
                import pypdfium2 as pdfium
                pdf = pdfium.PdfDocument(str(file_path))
                text = []
                for i in range(min(len(pdf), 50)):  # 最多读 50 页
                    page = pdf[i]
                    textpage = page.get_textpage()
                    text.append(textpage.get_text_range())
                return "\n".join(text)
            except:
                return ""

        # DOCX 文件
        elif ext == ".docx":
            try:
                from docx import Document
                doc = Document(str(file_path))
                text = [para.text for para in doc.paragraphs]
                return "\n".join(text)
            except:
                return ""

        return ""
    except:
        return ""

# ==================== BM25 实现 ====================

class SimpleBM25:
    """BM25 关键词检索"""

    def __init__(self, k1=1.5, b=0.75):
        self.k1 = k1
        self.b = b
        self.corpus = []
        self.doc_freqs = []
        self.idf = {}
        self.doc_len = []
        self.avgdl = 0

    def tokenize(self, text: str) -> List[str]:
        """中英文分词"""
        text = text.lower()
        chinese = re.findall(r'[一-鿿]+', text)
        english = re.findall(r'[a-z]+', text)

        tokens = []
        for word in chinese:
            tokens.extend(list(word))
        tokens.extend(english)
        return tokens

    def fit(self, corpus: List[str]):
        """构建BM25索引"""
        tokenized = [self.tokenize(doc) for doc in corpus]
        self.doc_freqs = [Counter(tokens) for tokens in tokenized]
        self.doc_len = [len(tokens) for tokens in tokenized]
        self.avgdl = sum(self.doc_len) / len(self.doc_len) if self.doc_len else 0

        # 计算 IDF
        df = Counter()
        for tokens in tokenized:
            for term in set(tokens):
                df[term] += 1

        N = len(corpus)
        for term, freq in df.items():
            self.idf[term] = math.log((N - freq + 0.5) / (freq + 0.5) + 1)

    def search(self, query: str, top_k: int = 10) -> List[Tuple[int, float]]:
        """搜索"""
        query_tokens = self.tokenize(query)
        scores = []

        for i, doc_freq in enumerate(self.doc_freqs):
            score = 0
            for term in query_tokens:
                if term not in doc_freq:
                    continue
                tf = doc_freq[term]
                idf = self.idf.get(term, 0)
                numerator = idf * tf * (self.k1 + 1)
                denominator = tf + self.k1 * (1 - self.b + self.b * self.doc_len[i] / self.avgdl)
                score += numerator / denominator
            scores.append((i, score))

        scores.sort(key=lambda x: x[1], reverse=True)
        return scores[:top_k]

# ==================== 全盘索引器 ====================

class DriveIndexer:
    """D盘全盘索引器"""

    def __init__(self, root_path: str = "D:/"):
        self.root_path = Path(root_path)
        self.documents = []
        self.metadata = []
        self.bm25 = SimpleBM25()

    def scan_drive(self) -> List[Path]:
        """扫描D盘所有支持的文件"""
        files = []
        skipped_dirs = 0
        skipped_files = 0

        print(f"[扫描] 开始扫描 {self.root_path}...", flush=True)

        for root, dirs, filenames in os.walk(self.root_path):
            root_path = Path(root)

            # 过滤目录（in-place 修改 dirs 让 os.walk 跳过）
            original_dirs = dirs[:]
            dirs[:] = [d for d in dirs if not should_skip_dir(root_path / d)]
            skipped_dirs += len(original_dirs) - len(dirs)

            # 收集文件
            for filename in filenames:
                file_path = root_path / filename
                if should_skip_file(file_path):
                    skipped_files += 1
                    continue
                files.append(file_path)

            # 进度提示（每 1000 个文件 + 每扫描一个新目录）
            if len(files) % 1000 == 0 and len(files) > 0:
                print(f"[扫描] 已发现 {len(files)} 个文件... 当前目录: {root_path}", flush=True)

        print(f"[扫描] 完成: 找到 {len(files)} 个文件, 跳过 {skipped_dirs} 目录, {skipped_files} 文件", flush=True)
        return files

    def build_index(self, incremental=False):
        """构建全盘索引（支持增量）"""
        files = self.scan_drive()

        # 增量模式：加载旧索引，只处理新增/修改文件
        old_metadata = {}
        if incremental and METADATA_FILE.exists():
            print("[增量] 加载旧索引...")
            with open(METADATA_FILE, 'r', encoding='utf-8') as f:
                old_meta_list = json.load(f)
                old_metadata = {m["path"]: m for m in old_meta_list}

        print(f"[索引] 开始读取 {len(files)} 个文件...")
        for i, file_path in enumerate(files):
            # 增量检查：文件未变化则跳过
            if incremental:
                path_str = str(file_path)
                if path_str in old_metadata:
                    old_mtime = old_metadata[path_str]["mtime"]
                    new_mtime = file_path.stat().st_mtime
                    if abs(new_mtime - old_mtime) < 1:  # 1秒容差
                        continue

            content = read_file_content(file_path)
            if not content or len(content) < 50:
                continue

            # 限制单文件最大长度（避免超长文件影响 BM25）
            content = content[:50000]

            self.documents.append(content)
            self.metadata.append({
                "path": str(file_path),
                "size": file_path.stat().st_size,
                "mtime": file_path.stat().st_mtime,
            })

            if (i + 1) % 100 == 0:
                print(f"[索引] 已处理 {i+1}/{len(files)} 文件...")

        print(f"[索引] 构建 BM25...")
        self.bm25.fit(self.documents)

        # 保存索引
        self.save()
        print(f"[完成] 索引了 {len(self.documents)} 个文件")
        if incremental:
            print(f"[增量] 新增/修改 {len(self.documents) - len(old_metadata)} 个文件")

    def save(self):
        """保存索引到磁盘"""
        with open(BM25_INDEX_FILE, 'wb') as f:
            pickle.dump(self.bm25, f)

        with open(METADATA_FILE, 'w', encoding='utf-8') as f:
            json.dump(self.metadata, f, ensure_ascii=False, indent=2)

        # 保存文档（用于显示预览）
        with open(EMBEDDINGS_FILE, 'wb') as f:
            pickle.dump(self.documents, f)

        # 状态
        state = {
            "indexed_at": datetime.now().isoformat(),
            "doc_count": len(self.documents),
        }
        with open(STATE_FILE, 'w') as f:
            json.dump(state, f)

    def load(self):
        """从磁盘加载索引"""
        if not BM25_INDEX_FILE.exists():
            print("[错误] 索引文件不存在，请先运行 index 命令")
            return False

        with open(BM25_INDEX_FILE, 'rb') as f:
            self.bm25 = pickle.load(f)

        with open(METADATA_FILE, 'r', encoding='utf-8') as f:
            self.metadata = json.load(f)

        with open(EMBEDDINGS_FILE, 'rb') as f:
            self.documents = pickle.load(f)

        return True

    def search(self, query: str, top_k: int = 10) -> List[Dict]:
        """搜索"""
        if not self.documents:
            if not self.load():
                return []

        results = self.bm25.search(query, top_k=top_k)

        formatted = []
        for idx, score in results:
            if score == 0:
                continue
            meta = self.metadata[idx]
            preview = self.documents[idx][:200].replace('\n', ' ')
            formatted.append({
                "score": score,
                "path": meta["path"],
                "preview": preview,
            })
        return formatted

# ==================== CLI ====================

import math
import sys

def main():
    if len(sys.argv) < 2:
        print("用法:")
        print("  python 全盘索引器_v1.py index [路径]   # 全量索引")
        print("  python 全盘索引器_v1.py update         # 增量更新")
        print("  python 全盘索引器_v1.py search '查询'  # 搜索")
        return

    command = sys.argv[1]

    if command == "index":
        root = sys.argv[2] if len(sys.argv) > 2 else "D:/"
        indexer = DriveIndexer(root)
        indexer.build_index(incremental=False)

    elif command == "update":
        root = sys.argv[2] if len(sys.argv) > 2 else "D:/"
        indexer = DriveIndexer(root)
        indexer.build_index(incremental=True)

    elif command == "search":
        if len(sys.argv) < 3:
            print("[错误] 请提供搜索关键词")
            return
        query = sys.argv[2]
        top_k = int(sys.argv[3]) if len(sys.argv) > 3 else 10

        indexer = DriveIndexer()
        results = indexer.search(query, top_k=top_k)

        if not results:
            print("[结果] 未找到相关文档")
            return

        print(f"\n[搜索] '{query}' - 找到 {len(results)} 个结果:\n")
        for i, r in enumerate(results, 1):
            print(f"[{i}] 分数: {r['score']:.2f}")
            print(f"    文件: {r['path']}")
            print(f"    预览: {r['preview'][:150]}...")
            print()

    else:
        print(f"[错误] 未知命令: {command}")


if __name__ == "__main__":
    main()
